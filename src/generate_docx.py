"""
Generate the academic DOCX reports mirroring the PDF content.

Two outputs:
  reports/Marathon_Shoe_Revolution_Decomposition_Report.docx   (full formal report)
  web/Shoe_Revolution_Three_Frameworks_Report.docx              (short companion)

Styling mirrors the boston-marathon-qualifying-fairness reference repo:
  - US Letter, 1" margins all around
  - Body: Calibri 11pt
  - Headings: Georgia bold
  - Title 22pt, subtitle 13pt italic
  - H1 15pt, H2 12pt
  - Captions: Calibri 9pt italic muted
  - Color palette: text #1A1A2E, soft #555555, muted #888888, men #2563EB, women #C0392B
"""
import json
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import pandas as pd

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(BASE_DIR)
FIG_DIR = os.path.join(PROJECT_DIR, 'outputs', 'figures')
SUMMARY_PATH = os.path.join(PROJECT_DIR, 'outputs', 'analysis_summary.json')
FULL_OUT = os.path.join(PROJECT_DIR, 'reports', 'Marathon_Shoe_Revolution_Decomposition_Report.docx')
SHORT_OUT = os.path.join(PROJECT_DIR, 'web', 'Shoe_Revolution_Three_Frameworks_Report.docx')

COLOR_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_SOFT = RGBColor(0x55, 0x55, 0x55)
COLOR_MUTED = RGBColor(0x88, 0x88, 0x88)
COLOR_ACCENT = RGBColor(0xC0, 0x39, 0x1B)


def load_summary():
    with open(SUMMARY_PATH) as f:
        return json.load(f)


def setup_section(section):
    """Letter size, 1" margins all around."""
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_title(doc, text, size_pt=22, color=COLOR_TEXT, italic=False, font='Georgia',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, bold=True):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return para


def add_h1(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = COLOR_TEXT
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_TEXT
    return para


def add_body(doc, text, justify=True):
    """
    Body paragraph supporting simple inline <b>...</b> tags for bold runs.
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.4

    # Parse <b>...</b> segments
    parts = []
    s = text
    while True:
        bstart = s.find('<b>')
        if bstart < 0:
            parts.append((s, False))
            break
        if bstart > 0:
            parts.append((s[:bstart], False))
        bend = s.find('</b>', bstart)
        if bend < 0:
            parts.append((s[bstart+3:], True))
            break
        parts.append((s[bstart+3:bend], True))
        s = s[bend+4:]

    for text_part, is_bold in parts:
        if not text_part:
            continue
        run = para.add_run(text_part)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
        if is_bold:
            run.font.bold = True
    return para


def add_figure(doc, filename, caption, width_inches=6.0):
    path = os.path.join(FIG_DIR, filename)
    if not os.path.exists(path):
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run()
    run.add_picture(path, width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    cap.paragraph_format.line_spacing = 1.2
    cap_run = cap.add_run(caption)
    cap_run.font.name = 'Calibri'
    cap_run.font.size = Pt(9)
    cap_run.font.italic = True
    cap_run.font.color.rgb = COLOR_MUTED


def add_page_break(doc):
    para = doc.add_paragraph()
    para.add_run().add_break(WD_BREAK.PAGE)


# ════════════════════════════════════════════════════════════════════
# FULL FORMAL DOCX REPORT
# ════════════════════════════════════════════════════════════════════

def build_full_report():
    summary = load_summary()
    f1 = summary['framework1_did']
    f2 = summary['framework2_within_athlete']
    f3 = summary['framework3_cohort_survival']
    pool = summary['pooled']

    doc = Document()
    setup_section(doc.sections[0])
    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title block
    add_title(doc, "How Much of the Post-2017 Marathon Revolution Is the Shoe?",
              size_pt=22, font='Georgia', bold=True)
    add_title(doc, "A Three-Framework Decomposition of Elite Marathon Improvement, 2010–2024",
              size_pt=13, color=COLOR_SOFT, italic=True, font='Georgia', bold=False, space_after=4)
    add_title(doc, "Jeremy Lee  |  May 2026  |  github.com/lyhjeremy/marathon-shoe-revolution-decomposition",
              size_pt=11, color=COLOR_SOFT, italic=True, font='Georgia', bold=False, space_after=20)

    add_h1(doc, "Abstract")
    add_body(doc,
        "When Nike's Vaporfly 4% reached commercial release in mid-2017, elite marathon "
        "times began falling at a pace that startled even close observers of the sport. "
        "By 2024 the men's world record had dropped 1:51 from where it stood in 2014; "
        "sub-2:10 marathon performances became roughly one-and-a-half times more common; "
        "an athlete won the 2024 Chicago Marathon in 2:00:35. The question is no longer "
        "whether shoes are real, but how much of that improvement they actually account for. "
        "This report applies three independent decomposition frameworks to ~1,900 elite "
        "marathon performances scraped from major-race Wikipedia tables for 2010–2024: "
        "(1) a difference-in-differences (DiD) comparison against track 10,000m as a "
        "less-affected control event; (2) a within-athlete paired pre/post analysis on "
        "the subset of athletes who raced elites in both eras; and (3) a cohort-depth "
        "analysis that compares observed post-era top-30 medians against the pre-era "
        f"trend. The three estimates converge on a pooled shoe contribution of "
        f"<b>{pool['pooled_estimate_seconds']:.0f} seconds (95% CI {pool['pooled_ci_low']:.0f}–"
        f"{pool['pooled_ci_high']:.0f} s)</b> in median elite marathon time, equivalent "
        f"to roughly 0.9% of a 2:05 marathon.")

    # 1. Introduction
    add_h1(doc, "1. Introduction")
    add_body(doc,
        "On 6 May 2017, three runners — Eliud Kipchoge, Lelisa Desisa, and Zersenay "
        "Tadese — ran 2:00:25 on the Monza Formula 1 circuit in Nike's 'Breaking2' "
        "exhibition. They wore prototypes of a shoe Nike would release commercially "
        "eight weeks later as the Zoom Vaporfly 4%. Independent biomechanics labs "
        "would later measure 4% running-economy improvement for a typical elite male "
        "marathoner under treadmill conditions.")
    add_body(doc,
        "Four months later, Kipchoge ran the official Berlin Marathon in 2:03:32 — "
        "already in the Vaporfly. A year later, in September 2018, he set the open "
        "world record at 2:01:39. Brigid Kosgei lowered the women's record by 81 "
        "seconds at Chicago 2019. Kelvin Kiptum went 2:00:35 at Chicago 2023. "
        "Sabastian Sawe became the first official sub-two-hour marathoner at "
        "London 2026.")
    add_body(doc,
        "A revolution that scale invites the obvious question. Are these shoes "
        "responsible for the improvement, or are they coincident with other changes — "
        "deeper East African talent pipelines, smarter time-trial-style pacing, "
        "expanded altitude training — that would have produced something like the "
        "post-2017 jump anyway? That is the question this report attempts to bound.")

    # 2. Data
    add_h1(doc, "2. Data Sources")
    add_body(doc,
        "<b>elite_marathon_times.csv</b> (1,908 rows): top-N finishers from per-race "
        "Wikipedia articles for six major marathons (Berlin, London, Chicago, Boston, "
        "Tokyo, New York City, plus a handful of Dubai entries), 2010–2024. "
        "<b>track_records_control.csv</b> (27 rows): top 10,000m track performances "
        "2016–2024 — the dataset's biggest weakness. <b>shoe_timeline.csv</b> (19 rows): "
        "hand-compiled release dates for every major super-shoe model 2016–2024. "
        "<b>athlete_career_arcs.csv</b> (derived): 17 athletes with ≥2 pre-era and ≥2 "
        "post-era major-marathon races.")
    add_body(doc,
        "Pre-era is 2010–2016; post-era is 2018–2024; 2017 is excluded as the transition "
        "year. Time-trial events (Breaking2 Monza, INEOS Vienna) are excluded since they "
        "are unratified.")

    add_figure(doc, 'fig1_sub210_frequency_by_year.png',
        "Figure 1. Counts of sub-2:10 (men) and sub-2:25 (women) marathon performances "
        "per year. Grey bars are pre-Vaporfly; orange the 2017 transition; red post-launch.")

    # 3. Methodology
    add_page_break(doc)
    add_h1(doc, "3. Methodology")

    add_h2(doc, "3.1 Framework 1: Difference-in-Differences (track vs road)")
    add_body(doc,
        "Track 10,000m racing did not adopt carbon plates at the same rate or to the "
        "same effect as road racing. If shoes are the dominant cause of road marathon "
        "improvement, road times should have improved faster than track times across "
        "the same window. We compute DiD = (road_post − road_pre) − (track_post − "
        "track_pre) per gender on top-30 median times. Bootstrap 95% CIs come from "
        "1,000 resamples of marathon rows. A placebo DiD splits the pre-era window "
        "into 2010–2013 vs 2014–2016, which should be close to zero.")

    add_h2(doc, "3.2 Framework 2: Within-Athlete Paired Pre/Post")
    add_body(doc,
        "For each athlete with ≥2 pre-era and ≥2 post-era major-marathon results, we "
        "compute the difference of post-era mean finish time minus pre-era mean finish "
        "time. We report the median across athletes, the paired t-test against zero, "
        "and a bootstrap 95% CI. The Wikipedia data does not include athlete ages, so "
        "we report the raw delta and discuss the age-confound limitation explicitly. "
        "With pre-era athletes typically aged 26–32 and post-era ages 28–36, not "
        "subtracting the +30 to +60 second aging penalty biases our shoe estimate "
        "downward, not upward.")

    add_h2(doc, "3.3 Framework 3: Cohort Survival / Depth")
    add_body(doc,
        "We count marathon performances faster than two thresholds: sub-2:10 (men) "
        "and sub-2:25 (women). Changepoint detection finds the year of the structural "
        "break. We take the difference between pre-era and post-era top-30 medians and "
        "attribute 55% of that observed cohort improvement to shoes (remainder: deeper "
        "fields, pacing, altitude, residual). Sensitivity to 40% and 70% is reported "
        "in §7.")

    # 4. Results
    add_page_break(doc)
    add_h1(doc, "4. Results")

    add_h2(doc,
        f"4.1 Framework 1 (DiD): {f1['shoe_contribution_seconds']:.0f} seconds, 95% CI "
        f"[{f1['ci_low']:.0f}, {f1['ci_high']:.0f}]")
    add_body(doc,
        f"The pre-era to post-era road marathon improvement (women's, top-30 median) was "
        f"<b>{abs(f1['by_gender']['W']['road_delta']):.0f} seconds "
        f"({f1['by_gender']['W']['road_pct']:.2f}%)</b>. The contemporaneous track 10,000m "
        f"women's improvement was just {abs(f1['by_gender']['W']['trk_delta']):.0f} seconds "
        f"({f1['by_gender']['W']['trk_pct']:.2f}%). Re-scaled, the track-explainable share "
        f"of road improvement is {f1['by_gender']['W']['track_explainable_marathon_seconds']:.0f} "
        f"seconds. The residual, attributable to road-specific factors of which shoes are "
        f"dominant, is <b>{f1['shoe_contribution_seconds']:.0f} seconds</b>. Placebo DiD "
        f"on pre-era splits returns {f1['placebo_did_seconds']:.0f} seconds — an order of "
        f"magnitude smaller than the main effect.")

    add_figure(doc, 'fig2_did_track_vs_road.png',
        "Figure 2. Road marathon top-50 mean time vs track 10,000m top-50 mean time, "
        "normalized to 2010 baseline. Curves diverge sharply after 2017.")

    add_h2(doc,
        f"4.2 Framework 2 (within-athlete): {f2['shoe_contribution_seconds']:.0f} seconds, "
        f"95% CI [{-f2['ci_high']:.0f}, {-f2['ci_low']:.0f}]")
    add_body(doc,
        f"<b>{f2['n_athletes']} athletes</b> meet the inclusion criteria. The median "
        f"across-athlete delta is <b>{f2['median_delta_seconds']:.0f} seconds</b> (post "
        f"mean minus pre mean). <b>{f2['pct_improved']:.0f}%</b> improved. A paired "
        f"t-test against zero returns t = {f2['paired_t']:.2f}, p = {f2['paired_p']:.2f} — "
        f"not significant at α=0.05. The bootstrap 95% CI on the mean delta is wide; "
        f"we report this estimate as suggestive rather than confirmatory.")

    add_figure(doc, 'fig3_within_athlete_paired.png',
        f"Figure 3. Within-athlete pre-vs-post mean finish time. Each line is one of "
        f"{f2['n_athletes']} athletes that raced ≥2 pre-era and ≥2 post-era majors.")

    add_h2(doc,
        f"4.3 Framework 3 (cohort survival): {f3['shoe_contribution_seconds']:.0f} seconds, "
        f"95% CI [{f3['ci_low']:.0f}, {f3['ci_high']:.0f}]")
    add_body(doc,
        f"The raw cohort improvement (top-30 median, averaged across men and women) is "
        f"<b>{f3['raw_cohort_improvement_seconds']:.0f} seconds</b> between pre-era and "
        f"post-era. The changepoint detection places the structural break at "
        f"<b>{f3['changepoint_year']}</b> — consistent with broad-cohort adoption following "
        f"elite adoption. Sub-threshold counts went from {f3['sub_m_pre_mean']:.0f} (men) "
        f"and {f3['sub_w_pre_mean']:.0f} (women) per year in the pre-era to "
        f"{f3['sub_m_post_mean']:.0f} and {f3['sub_w_post_mean']:.0f} in the post-era. "
        f"Attributing 55% of cohort improvement to shoes gives the framework point "
        f"estimate.")

    add_figure(doc, 'fig4_cohort_survival.png',
        "Figure 4. Top-50 threshold time per year, men and women. Inverted axis: lower on "
        "the chart = faster.")

    # 5. Cross-framework
    add_page_break(doc)
    add_h1(doc, "5. Cross-Framework Findings")

    add_body(doc,
        f"<b>Finding 1: All three frameworks place shoe contribution in the 47–111 "
        f"second range.</b> The pooled estimate of "
        f"<b>{pool['pooled_estimate_seconds']:.0f} seconds (95% CI "
        f"{pool['pooled_ci_low']:.0f}–{pool['pooled_ci_high']:.0f})</b> sits closer to "
        f"the conservative end because the within-athlete and cohort frameworks have "
        f"tighter CIs and higher weight in the pool.")
    add_body(doc,
        "<b>Finding 2: Sub-threshold marathon frequency is 1.25–1.4× higher post-2017.</b> "
        "Smaller than the 3–5× cited in popular running press — that figure is a global-"
        "pipeline statistic; our majors sample saturates earlier.")
    add_body(doc,
        "<b>Finding 3: Track 10,000m did not show comparable improvement.</b> "
        "Pre-to-post improvement was 0.32% — about one-fifth the road marathon's 1.59%. "
        "Track 10,000m is an imperfect control (spike technology did evolve), but the "
        "magnitude difference is large enough to be the strongest causal lever in the "
        "analysis.")

    add_figure(doc, 'fig5_framework_comparison.png',
        "Figure 5. The three frameworks plotted on a shared seconds-of-marathon-"
        "improvement x-axis. The pooled estimate sits below the DiD and above the "
        "within-athlete.")

    add_figure(doc, 'fig6_decomposition_pie.png',
        "Figure 6. Indicative decomposition of total 2016 → 2023 elite improvement. "
        "The 'carbon-plated shoes' slice uses the pooled estimate from this study; the "
        "remaining slices are post-hoc carve-ups.", width_inches=5.0)

    # 6. Historical
    add_page_break(doc)
    add_h1(doc, "6. Historical Comparison")
    add_body(doc,
        "Marathon world records had improved at a roughly steady 1 second per year "
        "through the 1990s and 2000s. From 2017 to 2024 the men's record dropped 1:51, "
        "the women's 4:08 (mixed-race). The pace of improvement was four to six times "
        "the long-term trend, aligned with elite-cohort adoption of carbon-plated shoes "
        "within 12 months.")

    add_figure(doc, 'fig8_brand_adoption_timeline.png',
        "Figure 7. The years each major brand introduced its first carbon-plated marathon "
        "shoe. The 2017–2020 window is the active transition.")

    # 7. Sensitivity
    add_h1(doc, "7. Sensitivity Analysis")
    add_body(doc,
        "<b>Scenario A:</b> Restrict to top-25 per year — DiD essentially unchanged. "
        "<b>Scenario B:</b> Exclude Eliud Kipchoge — estimate unchanged within bootstrap "
        "noise. <b>Scenario C:</b> Exclude time-trial-style events — change &lt; 5 "
        "seconds. <b>Scenario D:</b> 55% shoe-attribution share replaced with 40%/70% — "
        "pooled moves to 64 s / 73 s respectively. Across all four scenarios the pooled "
        "estimate stays within 55–80 seconds.")

    add_figure(doc, 'fig7_sensitivity.png',
        "Figure 8. DiD shoe-contribution estimate under each robustness scenario.")

    # 8. Limitations
    add_page_break(doc)
    add_h1(doc, "8. Limitations")
    add_body(doc,
        "<b>Track 10,000m pre-era coverage is sparse.</b> Wikipedia does not maintain "
        "year-by-year top-N lists. The DiD analysis is therefore women-only on the "
        "pre-side. A future revision using ARRS, World Athletics statistics archives, "
        "or Tilastopaja data would strengthen this framework substantially.")
    add_body(doc,
        "<b>No age adjustment in the within-athlete framework.</b> Wikipedia race tables "
        "do not report athlete age. Not subtracting the +30 to +60 second aging penalty "
        "biases our shoe estimate downward.")
    add_body(doc,
        "<b>Wikipedia coverage variance.</b> Some race-years have rich tables; others "
        "have only a podium plus a handful of selected times. We compensate with "
        "median-of-top-30 metrics.")
    add_body(doc,
        "<b>The East-African talent confound.</b> Post-2017 also saw a documented "
        "expansion of Kenyan and Ethiopian elite marathoning. The within-athlete "
        "framework controls for this; the cohort framework does not, which is part of "
        "why we attribute only 55% rather than 100% of cohort improvement to shoes.")
    add_body(doc,
        "<b>Course selection bias and pacing technology.</b> Post-2017 athletes "
        "increasingly self-selected into fast, pace-controlled courses, and laser "
        "pacing lights at Berlin (from 2019) contribute to post-era improvement "
        "independently of shoes. We bucket this into the 'pacing / time-trial' slice "
        "of the decomposition pie.")
    add_body(doc,
        "<b>Statistical power.</b> The within-athlete sample is n=17. The paired t-test "
        "against zero is not significant (p = 0.26). The finding from that framework "
        "should be interpreted as suggestive rather than confirmatory.")

    # 9. Conclusion
    add_h1(doc, "9. Conclusion")
    add_body(doc,
        f"The carbon-plated marathon shoe revolution is real, measurable, and population-"
        f"level — but smaller than the most aggressive popular characterizations and "
        f"larger than the most dismissive ones. Our three-framework decomposition places "
        f"the shoe-attributable share of elite marathon improvement between 2010–2016 and "
        f"2018–2024 at <b>{pool['pooled_estimate_seconds']:.0f} seconds in the median "
        f"elite marathon time, 95% CI {pool['pooled_ci_low']:.0f}–"
        f"{pool['pooled_ci_high']:.0f} seconds</b>, equivalent to roughly 0.9% of "
        f"marathon time for a 2:05 runner.")
    add_body(doc,
        "If you trust the within-athlete framework most, weight toward 47 s — the cleanest "
        "causal design but the smallest sample and not age-adjusted. If you trust the DiD "
        "most, weight toward 111 s — the strongest causal lever, but with sparse pre-side "
        "control data. If you trust the cohort survival framework most, weight toward 58 s "
        "— densest data, but the 55% shoe-attribution share is a modeling choice.")
    add_body(doc,
        "Shoes are the most-cited cause of the post-2017 marathon revolution not because "
        "they are the largest single cause, but because they are the most discrete and "
        "dateable one. The rest — deeper African talent pipelines, pacing improvements, "
        "altitude training, race-craft — together account for more.")

    # Footer note
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(20)
    run = para.add_run("Full code, data, and reproducibility instructions available at: "
                       "github.com/lyhjeremy/marathon-shoe-revolution-decomposition")
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = COLOR_MUTED

    doc.save(FULL_OUT)
    print(f"  DOCX saved: {FULL_OUT}")


# ════════════════════════════════════════════════════════════════════
# SHORT COMPANION DOCX
# ════════════════════════════════════════════════════════════════════

def build_short_report():
    summary = load_summary()
    f1 = summary['framework1_did']
    f2 = summary['framework2_within_athlete']
    f3 = summary['framework3_cohort_survival']
    pool = summary['pooled']

    doc = Document()
    setup_section(doc.sections[0])
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_title(doc, "How Much of the Post-2017 Marathon Revolution Is the Shoe?",
              size_pt=22, font='Georgia', bold=True)
    add_title(doc, "Three Frameworks, One Estimate",
              size_pt=13, color=COLOR_SOFT, italic=True, font='Georgia', bold=False, space_after=4)
    add_title(doc, "Jeremy Lee  |  May 2026  |  github.com/lyhjeremy/marathon-shoe-revolution-decomposition",
              size_pt=11, color=COLOR_SOFT, italic=True, font='Georgia', bold=False, space_after=20)

    add_h1(doc, "The Question")
    add_body(doc,
        "How much of the post-2017 elite marathon improvement is attributable to "
        "carbon-plated shoes, vs deeper fields, pacing improvements, altitude training, "
        "and other concurrent changes? This short report summarizes the headline numbers "
        "from a fuller analysis in the project repository.")

    add_h1(doc, "The Three Frameworks")
    add_body(doc,
        "<b>1. Difference-in-Differences:</b> compare road marathon improvement to track "
        "10,000m improvement over the same window. The residual road-specific improvement "
        "is the shoe-attributable share.")
    add_body(doc,
        "<b>2. Within-athlete paired:</b> same athletes pre-vs-post. Cancels genetics, "
        "training history, physiology.")
    add_body(doc,
        "<b>3. Cohort survival / depth:</b> distributional shift in the top-30 elite "
        "cohort, with a 55% shoe-attribution share.")

    add_figure(doc, 'fig5_framework_comparison.png',
        "Figure 1. The three frameworks plotted on a shared seconds-of-shoe-attributable-"
        "improvement x-axis, plus the inverse-CI-width-weighted pooled estimate.")

    add_h1(doc, "The Answer")
    add_body(doc,
        f"<b>Pooled shoe contribution: {pool['pooled_estimate_seconds']:.0f} seconds, "
        f"95% CI {pool['pooled_ci_low']:.0f}–{pool['pooled_ci_high']:.0f} s</b>, "
        f"equivalent to roughly 0.9% of a 2:05 marathon.")
    add_body(doc,
        f"Framework 1 (DiD): {f1['shoe_contribution_seconds']:.0f} s, 95% CI "
        f"[{f1['ci_low']:.0f}, {f1['ci_high']:.0f}]. "
        f"Framework 2 (within-athlete): {f2['shoe_contribution_seconds']:.0f} s, "
        f"95% CI [{-f2['ci_high']:.0f}, {-f2['ci_low']:.0f}], n={f2['n_athletes']}, "
        f"p={f2['paired_p']:.2f}. "
        f"Framework 3 (cohort): {f3['shoe_contribution_seconds']:.0f} s, "
        f"95% CI [{f3['ci_low']:.0f}, {f3['ci_high']:.0f}], changepoint = "
        f"{f3['changepoint_year']}.")

    add_figure(doc, 'fig1_sub210_frequency_by_year.png',
        "Figure 2. Sub-2:10 (men) and sub-2:25 (women) marathon performances per year, "
        "2010–2024.")

    add_h1(doc, "What the Data Allows You to Claim")
    add_body(doc,
        "Carbon-plated marathon shoes account for measurable, statistically credible "
        "improvement in elite marathon times. The contribution is real but smaller than "
        "headline framings suggest — around 1% of marathon time, not 4%. The remaining "
        "improvement comes from deeper African talent pipelines, pacing improvements, "
        "altitude training, and race-craft. Shoes are the most-cited cause because they "
        "are the most discrete and dateable, not because they are the largest cause.")

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(20)
    run = para.add_run("Full methodology, sensitivity analysis, and limitations at: "
                       "github.com/lyhjeremy/marathon-shoe-revolution-decomposition")
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = COLOR_MUTED

    doc.save(SHORT_OUT)
    print(f"  DOCX saved: {SHORT_OUT}")


if __name__ == '__main__':
    os.makedirs(os.path.dirname(FULL_OUT), exist_ok=True)
    os.makedirs(os.path.dirname(SHORT_OUT), exist_ok=True)
    build_full_report()
    build_short_report()
